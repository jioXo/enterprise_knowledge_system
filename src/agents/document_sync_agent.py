import asyncio
import hashlib
import json
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional
from pathlib import Path
import aiohttp
import fitz  # PyMuPDF
from docx import Document
import openpyxl

from src.models.database import SessionLocal
from src.models.document import Document, DocumentStatus, DocumentType
from src.models.platform import Platform, PlatformType
from src.models.knowledge import KnowledgeChunk, KnowledgeChunkStatus
from src.agents.base import BaseAgent
from src.utils.embedding import get_embedding_service
from src.utils.text_processor import TextProcessor
from src.config import get_settings

settings = get_settings()


class DocumentSyncAgent(BaseAgent):
    """文档同步与治理 Agent"""

    def __init__(self, config: Dict[str, Any]):
        super().__init__("DocumentSyncAgent", config)
        self.embedding_service = get_embedding_service()
        self.text_processor = TextProcessor()
        self.supported_formats = ['.txt', '.pdf', '.docx', '.xlsx', '.md']

    async def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """处理文档同步任务"""
        try:
            # 验证输入
            if not self.validate_input(input_data):
                return self.format_output({
                    "success": False,
                    "error": "Invalid input data"
                })

            platform_id = input_data.get("platform_id")
            force_sync = input_data.get("force_sync", False)

            # 获取数据库会话
            db = SessionLocal()

            try:
                # 获取平台信息
                platform = db.query(Platform).filter(Platform.id == platform_id).first()
                if not platform:
                    return self.format_output({
                        "success": False,
                        "error": f"Platform {platform_id} not found"
                    })

                # 执行同步
                if platform.platform_type == "local":
                    result = await self.sync_local_documents(db, platform, force_sync)
                elif platform.platform_type == "feishu":
                    result = await self.sync_feishu_documents(db, platform, force_sync)
                elif platform.platform_type == "confluence":
                    result = await self.sync_confluence_documents(db, platform, force_sync)
                else:
                    return self.format_output({
                        "success": False,
                        "error": f"Unsupported platform type: {platform.platform_type}"
                    })

                return self.format_output(result)

            finally:
                db.close()

        except Exception as e:
            self.log_error(e, "Document sync failed")
            return self.format_output({
                "success": False,
                "error": str(e)
            })

    def validate_input(self, input_data: Dict[str, Any]) -> bool:
        """验证输入数据"""
        required_fields = ["platform_id"]
        return all(field in input_data for field in required_fields)

    async def sync_local_documents(self, db: Session, platform: Platform, force_sync: bool) -> Dict[str, Any]:
        """同步本地文档"""
        self.log_info(f"Starting local document sync for platform {platform.name}")

        result = {
            "synced_count": 0,
            "updated_count": 0,
            "created_count": 0,
            "deleted_count": 0,
            "errors": []
        }

        # 扫描本地文档目录
        docs_path = Path(settings.RAW_DOCS_PATH) / str(platform.id)
        if not docs_path.exists():
            self.log_info(f"Documents directory not found: {docs_path}")
            return result

        # 遍历文档文件
        for doc_file in docs_path.rglob("*"):
            if doc_file.is_file() and doc_file.suffix.lower() in self.supported_formats:
                try:
                    doc_result = await self.process_document_file(db, platform, doc_file, force_sync)
                    result["synced_count"] += 1
                    result["created_count"] += doc_result["created"]
                    result["updated_count"] += doc_result["updated"]
                except Exception as e:
                    error_msg = f"Failed to process {doc_file}: {str(e)}"
                    self.log_error(e, error_msg)
                    result["errors"].append(error_msg)

        # 更新平台最后同步时间
        platform.last_sync_at = datetime.utcnow()
        db.commit()

        return result

    async def process_document_file(self, db: Session, platform: Platform, file_path: Path, force_sync: bool) -> Dict[str, Any]:
        """处理单个文档文件"""
        result = {"created": 0, "updated": 0}

        # 解析文档内容
        content, metadata = await self.parse_document_file(file_path)
        if not content:
            return result

        # 生成文档哈希（用于去重）
        content_hash = hashlib.md5(content.encode()).hexdigest()

        # 查找是否已存在文档
        existing_doc = db.query(Document).filter(
            Document.source_id == str(file_path),
            Document.platform_id == platform.id
        ).first()

        if existing_doc and not force_sync:
            # 检查内容是否有变化
            if existing_doc.content_hash == content_hash:
                return result

            # 更新现有文档
            await self.update_document(db, existing_doc, content, metadata, content_hash)
            result["updated"] = 1
        else:
            # 创建新文档
            await self.create_document(db, platform, content, metadata, content_hash, file_path)
            result["created"] = 1

        return result

    async def parse_document_file(self, file_path: Path) -> tuple:
        """解析文档文件"""
        try:
            if file_path.suffix.lower() == '.pdf':
                return await self.parse_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return await self.parse_docx(file_path)
            elif file_path.suffix.lower() == '.xlsx':
                return await self.parse_xlsx(file_path)
            elif file_path.suffix.lower() == '.txt':
                return await self.parse_txt(file_path)
            elif file_path.suffix.lower() == '.md':
                return await self.parse_markdown(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")

        except Exception as e:
            self.log_error(e, f"Failed to parse {file_path}")
            raise

    async def parse_pdf(self, file_path: Path) -> tuple:
        """解析PDF文件"""
        content = ""
        metadata = {"title": file_path.stem, "author": "", "word_count": 0}

        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    content += page.get_text() + "\n"

            # 提取元数据
            if doc.metadata:
                metadata["author"] = doc.metadata.get("author", "")
                metadata["title"] = doc.metadata.get("title", file_path.stem)

        except Exception as e:
            self.log_error(e, f"Failed to parse PDF {file_path}")
            return "", metadata

        # 处理文本
        content = self.text_processor.clean_text(content)
        metadata["word_count"] = len(content.split())

        return content, metadata

    async def parse_docx(self, file_path: Path) -> tuple:
        """解析Word文档"""
        try:
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            metadata = {
                "title": file_path.stem,
                "author": doc.core_properties.author if doc.core_properties else "",
                "word_count": len(content.split())
            }

            content = self.text_processor.clean_text(content)
            return content, metadata

        except Exception as e:
            self.log_error(e, f"Failed to parse DOCX {file_path}")
            return "", {"title": file_path.stem, "word_count": 0}

    async def parse_xlsx(self, file_path: Path) -> tuple:
        """解析Excel文件"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content = ""

            for sheet in workbook:
                content += f"## {sheet.title}\n"
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    content += row_text + "\n"
                content += "\n"

            metadata = {
                "title": file_path.stem,
                "author": workbook.properties.creator if workbook.properties else "",
                "word_count": len(content.split())
            }

            content = self.text_processor.clean_text(content)
            return content, metadata

        except Exception as e:
            self.log_error(e, f"Failed to parse XLSX {file_path}")
            return "", {"title": file_path.stem, "word_count": 0}

    async def parse_txt(self, file_path: Path) -> tuple:
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            metadata = {
                "title": file_path.stem,
                "word_count": len(content.split())
            }

            content = self.text_processor.clean_text(content)
            return content, metadata

        except Exception as e:
            self.log_error(e, f"Failed to parse TXT {file_path}")
            return "", {"title": file_path.stem, "word_count": 0}

    async def parse_markdown(self, file_path: Path) -> tuple:
        """解析Markdown文件"""
        return await self.parse_txt(file_path)  # 简单处理，后续可增强

    async def create_document(self, db: Session, platform: Platform, content: str, metadata: Dict, content_hash: str, file_path: Path):
        """创建新文档"""
        # 提取关键词
        keywords = self.text_processor.extract_keywords(content)
        summary = self.text_processor.generate_summary(content)

        document = Document(
            title=metadata.get("title", file_path.stem),
            content=content,
            summary=summary,
            document_type=self.determine_document_type(content),
            status=DocumentStatus.PUBLISHED,
            author=metadata.get("author", ""),
            department=self.extract_department(content),
            tags=",".join(keywords[:10]),  # 取前10个关键词
            version="1.0",
            word_count=metadata.get("word_count", 0),
            platform_id=platform.id,
            source_id=str(file_path),
            source_url=str(file_path.absolute()),
            content_hash=content_hash  # 添加哈希值
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        # 创建知识块
        await self.create_knowledge_chunks(db, document)

        self.log_info(f"Created new document: {document.title}")

    async def update_document(self, db: Session, document: Document, content: str, metadata: Dict, content_hash: str):
        """更新文档"""
        # 提取关键词
        keywords = self.text_processor.extract_keywords(content)
        summary = self.text_processor.generate_summary(content)

        # 更新文档内容
        document.content = content
        document.summary = summary
        document.document_type = self.determine_document_type(content)
        document.author = metadata.get("author", document.author)
        document.department = self.extract_department(content)
        document.tags = ",".join(keywords[:10])
        document.word_count = metadata.get("word_count", 0)
        document.content_hash = content_hash
        document.updated_at = datetime.utcnow()

        db.commit()

        # 更新知识块
        await self.update_knowledge_chunks(db, document)

        self.log_info(f"Updated document: {document.title}")

    def determine_document_type(self, content: str) -> DocumentType:
        """根据内容确定文档类型"""
        content_lower = content.lower()

        type_keywords = {
            DocumentType.POLICY: ["政策", "制度", "规定", "办法", "章程"],
            DocumentType.PROCEDURE: ["流程", "程序", "步骤", "操作", "指南"],
            DocumentType.FAQ: ["常见问题", "faq", "问题", "解答"],
            DocumentType.TRAINING: ["培训", "教程", "学习", "课程"],
            DocumentType.MANUAL: ["手册", "说明", "指南", "手册"],
            DocumentType.REPORT: ["报告", "总结", "分析", "汇报"],
        }

        type_scores = {}
        for doc_type, keywords in type_keywords.items():
            score = sum(content_lower.count(keyword) for keyword in keywords)
            type_scores[doc_type] = score

        # 返回得分最高的类型
        if type_scores:
            return max(type_scores, key=type_scores.get)
        return DocumentType.OTHER

    def extract_department(self, content: str) -> str:
        """从内容中提取部门信息"""
        # 简单实现，后续可使用NER增强
        department_keywords = ["人事", "人力资源", "财务", "行政", "技术", "研发", "市场", "销售", "运营"]
        content_lower = content.lower()

        for dept in department_keywords:
            if dept in content_lower:
                return dept
        return ""

    async def create_knowledge_chunks(self, db: Session, document: Document):
        """创建知识块"""
        # 文档分块
        chunks = self.text_processor.chunk_text(document.content, chunk_size=500, overlap=50)

        for i, chunk_content in enumerate(chunks):
            # 生成块摘要
            chunk_summary = self.text_processor.generate_summary(chunk_content)

            # 生成嵌入
            embedding = await self.embedding_service.generate_embedding(chunk_content)

            chunk = KnowledgeChunk(
                document_id=document.id,
                content=chunk_content,
                summary=chunk_summary,
                chunk_index=i,
                vector_id=f"doc_{document.id}_chunk_{i}",
                embedding=embedding.tolist() if embedding is not None else None,
                keywords=",".join(self.text_processor.extract_keywords(chunk_content)[:5]),
                topics=self.text_processor.extract_topics(chunk_content),
                importance_score=self.calculate_importance(chunk_content, document),
                word_count=len(chunk_content.split()),
                start_char=0,  # 简化实现
                end_char=len(chunk_content),
                status=KnowledgeChunkStatus.ACTIVE
            )

            db.add(chunk)

        db.commit()

    async def update_knowledge_chunks(self, db: Session, document: Document):
        """更新知识块"""
        # 获取现有知识块
        existing_chunks = db.query(KnowledgeChunk).filter(
            KnowledgeChunk.document_id == document.id
        ).all()

        # 删除过时的块（简化实现，实际应该做智能对比）
        for chunk in existing_chunks:
            db.delete(chunk)

        db.commit()

        # 重新创建知识块
        await self.create_knowledge_chunks(db, document)

    def calculate_importance(self, content: str, document: Document) -> float:
        """计算知识块的重要性分数"""
        # 简单实现：根据关键词、标题匹配度等计算
        score = 0.0

        # 标题附近的内容更重要
        if document.title.lower() in content.lower():
            score += 0.3

        # 包含关键词的内容更重要
        keywords = document.tags.split(",") if document.tags else []
        for keyword in keywords:
            if keyword.lower() in content.lower():
                score += 0.1

        # 长度适中的块更重要
        word_count = len(content.split())
        if 100 < word_count < 300:
            score += 0.2

        return min(score, 1.0)

    async def sync_feishu_documents(self, db: Session, platform: Platform, force_sync: bool) -> Dict[str, Any]:
        """同步飞书文档"""
        # TODO: 实现飞书文档同步
        self.log_info("Feishu document sync not implemented yet")
        return {"synced_count": 0, "updated_count": 0, "created_count": 0, "deleted_count": 0, "errors": []}

    async def sync_confluence_documents(self, db: Session, platform: Platform, force_sync: bool) -> Dict[str, Any]:
        """同步Confluence文档"""
        # TODO: 实现Confluence文档同步
        self.log_info("Confluence document sync not implemented yet")
        return {"synced_count": 0, "updated_count": 0, "created_count": 0, "deleted_count": 0, "errors": []}